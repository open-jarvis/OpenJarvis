"""Native Serena document operator tools.

Serena Documents Full Operator v1 foundation:
- locate/index documents
- read/extract supported file types
- summarize and classify
- inspect quality/completeness
- save reports locally
- protect originals from overwrite
"""

from __future__ import annotations

import json
import re
import shutil
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover
    DocxDocument = None

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None

from openjarvis.tools._stubs import BaseTool, ToolResult, ToolSpec
from openjarvis.core.registry import ToolRegistry


SUPPORTED_TEXT_SUFFIXES = {".txt", ".md", ".rtf"}
SUPPORTED_DOC_SUFFIXES = {".docx", ".pdf"}
SUPPORTED_SUFFIXES = SUPPORTED_TEXT_SUFFIXES | SUPPORTED_DOC_SUFFIXES


def _documents_root() -> Path:
    root = Path("outputs/documents")
    root.mkdir(parents=True, exist_ok=True)
    for child in ["library", "reports", "extracted", "summaries", "snapshots"]:
        (root / child).mkdir(parents=True, exist_ok=True)
    return root


def _safe_slug(value: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", value.strip().lower()).strip("-")
    return slug or "document"


def _timestamp() -> str:
    return time.strftime("%Y%m%d-%H%M%S")


def _resolve_path(path: str) -> Path:
    p = Path(path).expanduser()
    if not p.exists():
        raise FileNotFoundError(f"Document path not found: {p}")
    return p


def _read_text_like(path: Path) -> str:
    data = path.read_bytes()

    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue

    return data.decode("utf-8", errors="replace")


def _strip_rtf(text: str) -> str:
    # Lightweight RTF cleanup. Not a full RTF parser, but useful for many exported notes/docs.
    text = re.sub(r"{\\\*?\\[^{}]+}", " ", text)
    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", text)
    text = text.replace("{", " ").replace("}", " ")
    return re.sub(r"\s+", " ", text).strip()


def _extract_docx(path: Path) -> str:
    if DocxDocument is None:
        raise RuntimeError("python-docx is not installed. Run: uv add python-docx")

    try:
        doc = DocxDocument(str(path))
    except Exception as exc:
        raise RuntimeError(f"Could not open DOCX file. It may be corrupt, encrypted, or not a valid DOCX: {exc}") from exc

    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table_index, table in enumerate(doc.tables, start=1):
        parts.append(f"[Table {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    extracted = "\n".join(parts).strip()

    if not extracted:
        raise RuntimeError(
            "DOCX opened but no readable text was extracted. It may contain only images, shapes, or scanned content."
        )

    return extracted


def _extract_pdf(path: Path) -> str:
    if PdfReader is None:
        raise RuntimeError("pypdf is not installed. Run: uv add pypdf")

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise RuntimeError(f"Could not open PDF. It may be corrupt, encrypted, or password protected: {exc}") from exc

    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception as exc:
            raise RuntimeError(f"PDF is encrypted/password protected and could not be decrypted: {exc}") from exc

    page_texts: list[str] = []

    for index, page in enumerate(reader.pages, start=1):
        parts: list[str] = []

        try:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text.strip())
        except Exception as exc:
            parts.append(f"[Page {index}: text extraction failed: {exc}]")

        # Fallback: some PDFs contain useful text only in annotations/forms.
        try:
            annotations = page.get("/Annots") or []
            for annotation_ref in annotations:
                annotation = annotation_ref.get_object()
                contents = annotation.get("/Contents")
                if contents:
                    value = str(contents).strip()
                    if value:
                        parts.append(f"[Annotation] {value}")
        except Exception:
            pass

        if parts:
            page_texts.append(f"[Page {index}]\n" + "\n".join(parts))

    extracted = "\n\n".join(page_texts).strip()

    if not extracted:
        raise RuntimeError(
            "PDF opened but no readable text was extracted. This is likely a scanned/image-only PDF and needs OCR."
        )

    return extracted


def _extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return _read_text_like(path).strip()
    if suffix == ".rtf":
        return _strip_rtf(_read_text_like(path))
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    raise RuntimeError(f"Unsupported document type for v1 extraction: {suffix}")


def _word_count(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text))


def _line_count(text: str) -> int:
    return len([line for line in text.splitlines() if line.strip()])


def _classify_document(path: Path, text: str) -> dict[str, Any]:
    lower = text.lower()
    name = path.name.lower()

    scores = {
        "clinical_or_health": sum(k in lower for k in ["patient", "diagnosis", "treatment", "clinical", "medical", "health", "symptom", "doctor"]),
        "financial_or_billing": sum(k in lower for k in ["invoice", "payment", "claim", "billing", "amount", "tax", "medical aid", "statement"]),
        "legal_or_compliance": sum(k in lower for k in ["agreement", "contract", "terms", "privacy", "compliance", "policy", "consent", "liability"]),
        "marketing_or_content": sum(k in lower for k in ["seo", "blog", "newsletter", "campaign", "content", "social media", "cta", "landing page"]),
        "technical_or_project": sum(k in lower for k in ["api", "code", "system", "architecture", "deployment", "github", "server", "workflow"]),
        "general_document": 1,
    }

    if "cv" in name or "curriculum" in lower:
        scores["cv_or_profile"] = 5

    doc_type = max(scores, key=scores.get)
    confidence = min(0.95, 0.35 + (scores[doc_type] * 0.12))

    sensitivity_flags = []
    if scores.get("clinical_or_health", 0) >= 2:
        sensitivity_flags.append("healthcare/clinical")
    if scores.get("financial_or_billing", 0) >= 2:
        sensitivity_flags.append("financial/billing")
    if scores.get("legal_or_compliance", 0) >= 2:
        sensitivity_flags.append("legal/compliance")

    return {
        "document_type": doc_type,
        "confidence": round(confidence, 2),
        "scores": scores,
        "sensitivity_flags": sensitivity_flags,
    }


def _summarize_text(text: str, max_sentences: int = 6) -> str:
    clean = re.sub(r"\s+", " ", text).strip()
    if not clean:
        return "No readable text extracted."

    sentences = re.split(r"(?<=[.!?])\s+", clean)
    selected = [s.strip() for s in sentences if s.strip()][:max_sentences]

    if not selected:
        return clean[:1200]

    summary = " ".join(selected)
    return summary[:1800]


def _inspect_document(path: Path, text: str) -> dict[str, Any]:
    words = _word_count(text)
    lines = _line_count(text)

    issues = []
    recommendations = []

    suffix = path.suffix.lower()

    if not text.strip():
        issues.append("no_readable_text")
        recommendations.append("Use OCR or a different extraction path if this is scanned/image-based.")

    if suffix == ".pdf" and words < 20:
        issues.append("possible_scanned_pdf_or_ocr_needed")
        recommendations.append("PDF has very little extractable text. It may be scanned/image-only and need OCR.")

    if suffix == ".docx" and words < 20:
        issues.append("possible_empty_or_image_based_docx")
        recommendations.append("DOCX has very little extractable text. It may contain mostly images, shapes, or non-body text.")

    if words < 50:
        issues.append("very_short_text")
        recommendations.append("Confirm the document extracted correctly; it may be empty, scanned, or incomplete.")
    if words > 20000:
        recommendations.append("Large document detected. Consider section-based summaries.")
    if "signature" in text.lower() or "signed" in text.lower():
        recommendations.append("Signature-related language detected. Avoid altering original file.")
    if any(k in text.lower() for k in ["patient", "medical", "diagnosis", "treatment"]):
        recommendations.append("Healthcare content detected. Summarize and flag; do not make clinical decisions.")
    if any(k in text.lower() for k in ["contract", "agreement", "liability", "terms"]):
        recommendations.append("Legal/compliance language detected. Summarize and flag; do not provide final legal judgment.")

    return {
        "path": str(path),
        "filename": path.name,
        "suffix": path.suffix.lower(),
        "size_bytes": path.stat().st_size,
        "word_count": words,
        "line_count": lines,
        "issues": issues,
        "recommendations": recommendations,
    }


def _save_text_artifact(kind: str, source: Path, content: str, suffix: str = ".md") -> Path:
    root = _documents_root()
    out_dir = root / kind
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"{_timestamp()}-{_safe_slug(source.stem)}{suffix}"
    path.write_text(content, encoding="utf-8")
    return path


class _DocumentsBaseTool(BaseTool):
    def _result(self, content: str, success: bool = True, metadata: dict[str, Any] | None = None) -> ToolResult:
        return ToolResult(
            tool_name=getattr(self, "tool_id", self.__class__.__name__),
            success=success,
            content=content,
            metadata=metadata or {},
        )


@ToolRegistry.register("serena_documents_status")
class SerenaDocumentsStatusTool(_DocumentsBaseTool):
    tool_id = "serena_documents_status"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Show Serena Documents Full Operator v1 status and supported formats.",
            parameters={"type": "object", "properties": {}},
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        root = _documents_root()
        return self._result(
            "Serena Documents status\n\n"
            "- Status: active\n"
            "- Supported v1 formats: .txt, .md, .rtf, .docx, .pdf\n"
            "- PDF text extraction: active\n"
            "- OCR/Google Drive: planned next layers\n"
            f"- Output root: {root}\n"
            f"- Reports: {root / 'reports'}\n"
            f"- Extracted text: {root / 'extracted'}\n"
            f"- Summaries: {root / 'summaries'}\n"
            f"- Snapshots: {root / 'snapshots'}",
            metadata={"root": str(root), "supported_formats": sorted(SUPPORTED_SUFFIXES)},
        )


@ToolRegistry.register("serena_documents_index")
class SerenaDocumentsIndexTool(_DocumentsBaseTool):
    tool_id = "serena_documents_index"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Index supported local documents in a folder.",
            parameters={
                "type": "object",
                "properties": {
                    "folder": {"type": "string", "description": "Folder to scan."},
                    "recursive": {"type": "boolean", "description": "Scan recursively."},
                    "limit": {"type": "integer", "description": "Maximum files to list."},
                },
                "required": ["folder"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        folder = _resolve_path(str(params.get("folder") or ""))
        recursive = bool(params.get("recursive", True))
        limit = int(params.get("limit") or 100)

        if not folder.is_dir():
            return self._result(f"Not a folder: {folder}", success=False)

        pattern = "**/*" if recursive else "*"
        files = [
            p for p in folder.glob(pattern)
            if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES
        ][:limit]

        lines = [
            "Serena document index",
            "",
            f"- Folder: {folder}",
            f"- Recursive: {'yes' if recursive else 'no'}",
            f"- Supported documents found: {len(files)}",
            "",
            "Documents:",
        ]

        if not files:
            lines.append("- none")
        else:
            for p in files:
                lines.append(f"- {p.name} | {p.suffix.lower()} | {p.stat().st_size} bytes | {p}")

        index_path = _save_text_artifact("reports", folder, "\n".join(lines), suffix=".md")

        return self._result(
            "\n".join(lines) + f"\n\nIndex saved: {index_path}",
            metadata={"folder": str(folder), "count": len(files), "index_path": str(index_path)},
        )


@ToolRegistry.register("serena_documents_extract")
class SerenaDocumentsExtractTool(_DocumentsBaseTool):
    tool_id = "serena_documents_extract"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Extract readable text from a supported local document and save it.",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Document path."}
                },
                "required": ["path"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            path = _resolve_path(str(params.get("path") or ""))
            text = _extract_text(path)
            out_path = _save_text_artifact("extracted", path, text, suffix=".txt")
            inspection = _inspect_document(path, text)

            return self._result(
                "Serena document text extracted\n\n"
                f"- Source: {path}\n"
                f"- Output: {out_path}\n"
                f"- Word count: {inspection['word_count']}\n"
                f"- Line count: {inspection['line_count']}",
                metadata={"source": str(path), "output": str(out_path), **inspection},
            )
        except Exception as exc:
            return self._result(f"Failed to extract document text: {exc}", success=False)


@ToolRegistry.register("serena_documents_read")
class SerenaDocumentsReadTool(SerenaDocumentsExtractTool):
    tool_id = "serena_documents_read"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Read/extract a supported local document.",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Document path."},
                    "preview_chars": {"type": "integer", "description": "Preview character count."},
                },
                "required": ["path"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        result = super().execute(**params)
        if not result.success:
            return result

        output = Path(result.metadata["output"])
        text = output.read_text(encoding="utf-8")
        preview_chars = int(params.get("preview_chars") or 2000)
        preview = text[:preview_chars]

        return self._result(
            result.content + "\n\nPreview:\n" + preview,
            metadata={**result.metadata, "preview_chars": preview_chars},
        )


@ToolRegistry.register("serena_documents_summarize")
class SerenaDocumentsSummarizeTool(_DocumentsBaseTool):
    tool_id = "serena_documents_summarize"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Extract and summarize a supported local document.",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Document path."},
                    "max_sentences": {"type": "integer", "description": "Maximum summary sentences."},
                },
                "required": ["path"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            path = _resolve_path(str(params.get("path") or ""))
            text = _extract_text(path)
            summary = _summarize_text(text, max_sentences=int(params.get("max_sentences") or 6))
            classification = _classify_document(path, text)
            inspection = _inspect_document(path, text)

            report = (
                f"# Document Summary\n\n"
                f"Source: `{path}`\n\n"
                f"## Summary\n\n{summary}\n\n"
                f"## Classification\n\n"
                f"- Type: {classification['document_type']}\n"
                f"- Confidence: {classification['confidence']}\n"
                f"- Sensitivity flags: {', '.join(classification['sensitivity_flags']) or 'none'}\n\n"
                f"## Inspection\n\n"
                f"- Word count: {inspection['word_count']}\n"
                f"- Line count: {inspection['line_count']}\n"
                f"- Issues: {', '.join(inspection['issues']) or 'none'}\n"
            )

            out_path = _save_text_artifact("summaries", path, report, suffix=".md")

            return self._result(
                "Serena document summary created\n\n"
                f"- Source: {path}\n"
                f"- Summary: {out_path}\n"
                f"- Type: {classification['document_type']}\n"
                f"- Sensitivity flags: {', '.join(classification['sensitivity_flags']) or 'none'}\n\n"
                f"Summary preview:\n{summary}",
                metadata={"source": str(path), "summary_path": str(out_path), **classification, **inspection},
            )
        except Exception as exc:
            return self._result(f"Failed to summarize document: {exc}", success=False)


@ToolRegistry.register("serena_documents_classify")
class SerenaDocumentsClassifyTool(_DocumentsBaseTool):
    tool_id = "serena_documents_classify"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Classify a supported local document and flag sensitivity.",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Document path."}
                },
                "required": ["path"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            path = _resolve_path(str(params.get("path") or ""))
            text = _extract_text(path)
            classification = _classify_document(path, text)
            inspection = _inspect_document(path, text)

            return self._result(
                "Serena document classification\n\n"
                f"- Source: {path}\n"
                f"- Type: {classification['document_type']}\n"
                f"- Confidence: {classification['confidence']}\n"
                f"- Sensitivity flags: {', '.join(classification['sensitivity_flags']) or 'none'}\n"
                f"- Word count: {inspection['word_count']}",
                metadata={"source": str(path), **classification, **inspection},
            )
        except Exception as exc:
            return self._result(f"Failed to classify document: {exc}", success=False)


@ToolRegistry.register("serena_documents_inspect")
class SerenaDocumentsInspectTool(_DocumentsBaseTool):
    tool_id = "serena_documents_inspect"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Inspect a supported local document for extraction quality, completeness, and safety flags.",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Document path."}
                },
                "required": ["path"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            path = _resolve_path(str(params.get("path") or ""))

            try:
                text = _extract_text(path)
                extraction_error = ""
            except Exception as extract_exc:
                text = ""
                extraction_error = str(extract_exc)

            inspection = _inspect_document(path, text)
            classification = _classify_document(path, text)

            if extraction_error:
                inspection["issues"].append("text_extraction_failed")
                if path.suffix.lower() == ".pdf":
                    inspection["issues"].append("pdf_ocr_or_password_check_needed")
                    inspection["recommendations"].append("Run pdf-check. If no text is extractable, use OCR before summarizing.")
                inspection["recommendations"].append(f"Extraction error: {extraction_error}")

            lines = [
                "Serena document inspection",
                "",
                f"- Source: {path}",
                f"- File type: {inspection['suffix']}",
                f"- Size: {inspection['size_bytes']} bytes",
                f"- Word count: {inspection['word_count']}",
                f"- Line count: {inspection['line_count']}",
                f"- Classified as: {classification['document_type']} ({classification['confidence']})",
                f"- Sensitivity flags: {', '.join(classification['sensitivity_flags']) or 'none'}",
                "",
                "Issues:",
            ]

            lines.extend(f"- {issue}" for issue in inspection["issues"] or ["none"])
            lines.append("")
            lines.append("Recommendations:")
            lines.extend(f"- {rec}" for rec in inspection["recommendations"] or ["No major recommendations."])

            return self._result("\n".join(lines), metadata={**inspection, **classification, "extraction_error": extraction_error})
        except Exception as exc:
            return self._result(f"Failed to inspect document: {exc}", success=False)


@ToolRegistry.register("serena_documents_report")
class SerenaDocumentsReportTool(_DocumentsBaseTool):
    tool_id = "serena_documents_report"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Create a complete Serena document operator report for a supported local document.",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Document path."}
                },
                "required": ["path"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            path = _resolve_path(str(params.get("path") or ""))
            text = _extract_text(path)
            classification = _classify_document(path, text)
            inspection = _inspect_document(path, text)
            summary = _summarize_text(text, max_sentences=8)

            report = f"""# Serena Document Operator Report

Source: `{path}`

## Identity

- Filename: {path.name}
- Type: {path.suffix.lower()}
- Size: {path.stat().st_size} bytes

## Classification

- Document type: {classification['document_type']}
- Confidence: {classification['confidence']}
- Sensitivity flags: {', '.join(classification['sensitivity_flags']) or 'none'}

## Extraction Quality

- Word count: {inspection['word_count']}
- Line count: {inspection['line_count']}
- Issues: {', '.join(inspection['issues']) or 'none'}

## Summary

{summary}

## Recommendations

{chr(10).join('- ' + rec for rec in inspection['recommendations']) if inspection['recommendations'] else '- No major recommendations.'}

## Operator Notes

- Original file was not modified.
- Generated outputs are stored under `outputs/documents/`.
- For healthcare/legal/financial documents, Serena provides summaries and flags, not final professional decisions.
"""

            out_path = _save_text_artifact("reports", path, report, suffix=".md")

            return self._result(
                "Serena document operator report created\n\n"
                f"- Source: {path}\n"
                f"- Report: {out_path}\n"
                f"- Type: {classification['document_type']}\n"
                f"- Word count: {inspection['word_count']}",
                metadata={"source": str(path), "report_path": str(out_path), **classification, **inspection},
            )
        except Exception as exc:
            return self._result(f"Failed to create document report: {exc}", success=False)


def _library_dir(category: str = "general") -> Path:
    root = _documents_root() / "library" / _safe_slug(category)
    root.mkdir(parents=True, exist_ok=True)
    return root


def _snapshot_document(path: Path, reason: str) -> Path:
    root = _documents_root() / "snapshots"
    root.mkdir(parents=True, exist_ok=True)

    timestamp = _timestamp()
    target = root / f"{timestamp}-{_safe_slug(path.stem)}-{_safe_slug(reason)}{path.suffix.lower()}"

    if path.exists() and path.is_file():
        shutil.copy2(path, target)

    meta = {
        "source": str(path),
        "snapshot": str(target),
        "reason": reason,
        "timestamp": timestamp,
        "size_bytes": path.stat().st_size if path.exists() else 0,
    }

    meta_path = target.with_suffix(target.suffix + ".json")
    meta_path.write_text(json.dumps(meta, indent=2), encoding="utf-8")

    return target


@ToolRegistry.register("serena_documents_import")
class SerenaDocumentsImportTool(_DocumentsBaseTool):
    tool_id = "serena_documents_import"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Copy a document into Serena's controlled document library without modifying the original.",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Source document path."},
                    "category": {"type": "string", "description": "Library category/folder."},
                },
                "required": ["path"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            source = _resolve_path(str(params.get("path") or ""))
            category = str(params.get("category") or "general").strip() or "general"

            if not source.is_file():
                return self._result(f"Not a file: {source}", success=False)

            if source.suffix.lower() not in SUPPORTED_SUFFIXES:
                return self._result(
                    f"Unsupported v1 document type: {source.suffix.lower()}",
                    success=False,
                    metadata={"supported": sorted(SUPPORTED_SUFFIXES)},
                )

            library = _library_dir(category)
            target = library / f"{_timestamp()}-{_safe_slug(source.stem)}{source.suffix.lower()}"

            shutil.copy2(source, target)

            text = _extract_text(target)
            classification = _classify_document(target, text)
            inspection = _inspect_document(target, text)

            meta = {
                "source": str(source),
                "library_path": str(target),
                "category": category,
                "imported_at": _timestamp(),
                "classification": classification,
                "inspection": inspection,
                "original_preserved": True,
            }

            meta_path = target.with_suffix(target.suffix + ".json")
            meta_path.write_text(json.dumps(meta, indent=2), encoding="utf-8")

            return self._result(
                "Serena document imported into controlled library\n\n"
                f"- Source: {source}\n"
                f"- Library path: {target}\n"
                f"- Metadata: {meta_path}\n"
                f"- Category: {category}\n"
                f"- Type: {classification['document_type']}\n"
                f"- Sensitivity flags: {', '.join(classification['sensitivity_flags']) or 'none'}\n"
                f"- Original preserved: yes",
                metadata={
                    "source": str(source),
                    "library_path": str(target),
                    "metadata_path": str(meta_path),
                    "category": category,
                    **classification,
                    **inspection,
                },
            )
        except Exception as exc:
            return self._result(f"Failed to import document: {exc}", success=False)


@ToolRegistry.register("serena_documents_library")
class SerenaDocumentsLibraryTool(_DocumentsBaseTool):
    tool_id = "serena_documents_library"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="List Serena's controlled document library.",
            parameters={
                "type": "object",
                "properties": {
                    "category": {"type": "string", "description": "Optional category/folder."},
                    "limit": {"type": "integer", "description": "Maximum files to show."},
                },
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        category = str(params.get("category") or "").strip()
        limit = int(params.get("limit") or 50)

        root = _documents_root() / "library"
        folder = _library_dir(category) if category else root

        files = [
            p for p in folder.glob("**/*")
            if p.is_file()
            and p.suffix.lower() in SUPPORTED_SUFFIXES
        ]

        files = sorted(files, key=lambda p: p.stat().st_mtime, reverse=True)[:limit]

        lines = [
            "Serena document library",
            "",
            f"- Folder: {folder}",
            f"- Documents found: {len(files)}",
            "",
            "Documents:",
        ]

        if not files:
            lines.append("- none")
        else:
            for file in files:
                lines.append(f"- {file.name} | {file.suffix.lower()} | {file.stat().st_size} bytes | {file}")

        return self._result(
            "\n".join(lines),
            metadata={"folder": str(folder), "count": len(files)},
        )


@ToolRegistry.register("serena_documents_snapshot")
class SerenaDocumentsSnapshotTool(_DocumentsBaseTool):
    tool_id = "serena_documents_snapshot"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Create a local safety snapshot of a document before any risky operation.",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Document path."},
                    "reason": {"type": "string", "description": "Snapshot reason."},
                },
                "required": ["path"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            source = _resolve_path(str(params.get("path") or ""))
            reason = str(params.get("reason") or "manual-snapshot").strip() or "manual-snapshot"

            if not source.is_file():
                return self._result(f"Not a file: {source}", success=False)

            snapshot = _snapshot_document(source, reason)

            return self._result(
                "Serena document snapshot created\n\n"
                f"- Source: {source}\n"
                f"- Snapshot: {snapshot}\n"
                f"- Reason: {reason}",
                metadata={"source": str(source), "snapshot": str(snapshot), "reason": reason},
            )
        except Exception as exc:
            return self._result(f"Failed to snapshot document: {exc}", success=False)


@ToolRegistry.register("serena_documents_snapshots")
class SerenaDocumentsSnapshotsTool(_DocumentsBaseTool):
    tool_id = "serena_documents_snapshots"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="List Serena document snapshots.",
            parameters={
                "type": "object",
                "properties": {
                    "limit": {"type": "integer", "description": "Maximum snapshots to show."},
                },
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        limit = int(params.get("limit") or 50)
        folder = _documents_root() / "snapshots"

        files = [
            p for p in folder.glob("*")
            if p.is_file()
            and not p.name.endswith(".json")
        ]

        files = sorted(files, key=lambda p: p.stat().st_mtime, reverse=True)[:limit]

        lines = [
            "Serena document snapshots",
            "",
            f"- Folder: {folder}",
            f"- Snapshots found: {len(files)}",
            "",
            "Snapshots:",
        ]

        if not files:
            lines.append("- none")
        else:
            for file in files:
                lines.append(f"- {file.name} | {file.stat().st_size} bytes | {file}")

        return self._result(
            "\n".join(lines),
            metadata={"folder": str(folder), "count": len(files)},
        )


@ToolRegistry.register("serena_documents_audit")
class SerenaDocumentsAuditTool(_DocumentsBaseTool):
    tool_id = "serena_documents_audit"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Run Serena's document library audit dashboard.",
            parameters={
                "type": "object",
                "properties": {
                    "folder": {"type": "string", "description": "Optional folder to audit. Defaults to outputs/documents/library."},
                    "recursive": {"type": "boolean", "description": "Scan recursively."},
                    "limit": {"type": "integer", "description": "Maximum files to audit."},
                },
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            root = _documents_root()
            folder_value = str(params.get("folder") or "").strip()
            folder = Path(folder_value) if folder_value else root / "library"
            recursive = bool(params.get("recursive", True))
            limit = int(params.get("limit") or 200)

            if not folder.exists() or not folder.is_dir():
                return self._result(f"Audit folder not found or not a folder: {folder}", success=False)

            pattern = "**/*" if recursive else "*"
            all_files = [p for p in folder.glob(pattern) if p.is_file()]
            document_files = [p for p in all_files if p.suffix.lower() in SUPPORTED_SUFFIXES][:limit]
            unsupported_files = [
                p for p in all_files
                if not p.name.endswith(".json")
                and p.suffix.lower() not in SUPPORTED_SUFFIXES
            ][:limit]

            audited: list[dict[str, Any]] = []
            extraction_failures: list[dict[str, str]] = []

            for doc_path in document_files:
                try:
                    text = _extract_text(doc_path)
                    classification = _classify_document(doc_path, text)
                    inspection = _inspect_document(doc_path, text)
                    meta_path = doc_path.with_suffix(doc_path.suffix + ".json")
                    audited.append(
                        {
                            "path": str(doc_path),
                            "name": doc_path.name,
                            "category": str(doc_path.parent.relative_to(root / "library")) if (root / "library") in doc_path.parents else "external",
                            "type": classification["document_type"],
                            "confidence": classification["confidence"],
                            "sensitivity_flags": classification["sensitivity_flags"],
                            "word_count": inspection["word_count"],
                            "line_count": inspection["line_count"],
                            "issues": inspection["issues"],
                            "recommendations": inspection["recommendations"],
                            "metadata_exists": meta_path.exists(),
                        }
                    )
                except Exception as exc:
                    extraction_failures.append({"path": str(doc_path), "error": str(exc)})

            sensitive = [x for x in audited if x["sensitivity_flags"]]
            short_or_empty = [x for x in audited if x["word_count"] < 50]
            missing_metadata = [x for x in audited if not x["metadata_exists"]]

            categories: dict[str, int] = {}
            doc_types: dict[str, int] = {}
            for item in audited:
                categories[item["category"]] = categories.get(item["category"], 0) + 1
                doc_types[item["type"]] = doc_types.get(item["type"], 0) + 1

            recent_reports = sorted((root / "reports").glob("*.md"), key=lambda p: p.stat().st_mtime, reverse=True)[:10]
            recent_summaries = sorted((root / "summaries").glob("*.md"), key=lambda p: p.stat().st_mtime, reverse=True)[:10]
            recent_snapshots = sorted(
                [p for p in (root / "snapshots").glob("*") if p.is_file() and not p.name.endswith(".json")],
                key=lambda p: p.stat().st_mtime,
                reverse=True,
            )[:10]

            lines = [
                "Serena Documents audit dashboard",
                "",
                f"- Audit folder: {folder}",
                f"- Recursive: {'yes' if recursive else 'no'}",
                f"- Supported documents scanned: {len(audited)}",
                f"- Unsupported files found: {len(unsupported_files)}",
                f"- Extraction failures: {len(extraction_failures)}",
                "",
                "Summary:",
                f"- Sensitive documents flagged: {len(sensitive)}",
                f"- Very short/empty documents: {len(short_or_empty)}",
                f"- Documents missing metadata: {len(missing_metadata)}",
                f"- Recent reports: {len(recent_reports)}",
                f"- Recent summaries: {len(recent_summaries)}",
                f"- Recent snapshots: {len(recent_snapshots)}",
                "",
                "Documents by category:",
            ]

            if categories:
                for category, count in sorted(categories.items()):
                    lines.append(f"- {category}: {count}")
            else:
                lines.append("- none")

            lines.extend(["", "Documents by type:"])
            if doc_types:
                for doc_type, count in sorted(doc_types.items()):
                    lines.append(f"- {doc_type}: {count}")
            else:
                lines.append("- none")

            lines.extend(["", "Sensitive documents:"])
            if sensitive:
                for item in sensitive[:15]:
                    lines.append(
                        f"- {item['name']} | {item['type']} | flags: {', '.join(item['sensitivity_flags'])} | words {item['word_count']}"
                    )
            else:
                lines.append("- none")

            lines.extend(["", "Very short or empty documents:"])
            if short_or_empty:
                for item in short_or_empty[:15]:
                    lines.append(f"- {item['name']} | words {item['word_count']} | issues: {', '.join(item['issues']) or 'none'}")
            else:
                lines.append("- none")

            lines.extend(["", "Documents missing metadata:"])
            if missing_metadata:
                for item in missing_metadata[:15]:
                    lines.append(f"- {item['name']} | {item['path']}")
            else:
                lines.append("- none")

            lines.extend(["", "Unsupported files:"])
            if unsupported_files:
                for item in unsupported_files[:15]:
                    lines.append(f"- {item.name} | {item.suffix.lower()} | {item}")
            else:
                lines.append("- none")

            lines.extend(["", "Extraction failures:"])
            if extraction_failures:
                for item in extraction_failures[:15]:
                    lines.append(f"- {item['path']}: {item['error']}")
            else:
                lines.append("- none")

            lines.extend(["", "Recent reports:"])
            lines.extend(f"- {p.name}" for p in recent_reports) if recent_reports else lines.append("- none")

            lines.extend(["", "Recent summaries:"])
            lines.extend(f"- {p.name}" for p in recent_summaries) if recent_summaries else lines.append("- none")

            lines.extend(["", "Recent snapshots:"])
            lines.extend(f"- {p.name}" for p in recent_snapshots) if recent_snapshots else lines.append("- none")

            lines.extend([
                "",
                "Recommended operator actions:",
                "- Add metadata by importing unmanaged supported documents into Serena's document library.",
                "- Review sensitive documents carefully; summarize and flag, but do not make clinical/legal/financial decisions.",
                "- Use OCR/PDF layer later for scanned or unsupported documents.",
                "- Create snapshots before any move, rename, or cleanup action.",
                "- Generate reports for important documents before acting on them.",
            ])

            audit_report = _save_text_artifact("reports", folder, "\n".join(lines), suffix=".md")

            return self._result(
                "\n".join(lines) + f"\n\nAudit report saved: {audit_report}",
                metadata={
                    "folder": str(folder),
                    "documents_scanned": len(audited),
                    "unsupported_files": len(unsupported_files),
                    "extraction_failures": len(extraction_failures),
                    "sensitive_documents": len(sensitive),
                    "short_or_empty": len(short_or_empty),
                    "missing_metadata": len(missing_metadata),
                    "audit_report": str(audit_report),
                },
            )
        except Exception as exc:
            return self._result(f"Failed to run document audit: {exc}", success=False)


@ToolRegistry.register("serena_documents_pdf_check")
class SerenaDocumentsPDFCheckTool(_DocumentsBaseTool):
    tool_id = "serena_documents_pdf_check"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Check a PDF for readable text extraction and flag likely scanned/OCR-needed PDFs.",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "PDF path."}
                },
                "required": ["path"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            path = _resolve_path(str(params.get("path") or ""))

            if path.suffix.lower() != ".pdf":
                return self._result(f"Not a PDF file: {path}", success=False)

            if PdfReader is None:
                return self._result("pypdf is not installed. Run: uv add pypdf", success=False)

            reader = PdfReader(str(path))
            encrypted = bool(getattr(reader, "is_encrypted", False))
            pages = len(reader.pages)

            extracted_pages = 0
            failed_pages = 0
            total_words = 0

            if encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    pass

            for page in reader.pages:
                try:
                    text = page.extract_text() or ""
                    words = _word_count(text)
                    total_words += words
                    if words > 0:
                        extracted_pages += 1
                except Exception:
                    failed_pages += 1

            likely_scanned = total_words < 20 or extracted_pages == 0

            lines = [
                "Serena PDF extraction check",
                "",
                f"- Source: {path}",
                f"- Pages: {pages}",
                f"- Encrypted: {'yes' if encrypted else 'no'}",
                f"- Pages with extractable text: {extracted_pages}",
                f"- Failed pages: {failed_pages}",
                f"- Extracted word count estimate: {total_words}",
                f"- Likely scanned/OCR-needed: {'yes' if likely_scanned else 'no'}",
                "",
                "Recommendation:",
            ]

            if likely_scanned:
                lines.append("- Use the future OCR layer for this PDF, or convert it with OCR before summarizing.")
            else:
                lines.append("- PDF has extractable text and can be summarized/reported by Serena Documents.")

            return self._result(
                "\n".join(lines),
                metadata={
                    "source": str(path),
                    "pages": pages,
                    "encrypted": encrypted,
                    "pages_with_text": extracted_pages,
                    "failed_pages": failed_pages,
                    "word_count_estimate": total_words,
                    "likely_scanned": likely_scanned,
                },
            )
        except Exception as exc:
            return self._result(f"Failed to check PDF: {exc}", success=False)


def _extract_structured_fields(path: Path, text: str) -> dict[str, Any]:
    """Extract practical structured fields from document text using safe deterministic patterns."""
    clean = re.sub(r"\s+", " ", text).strip()
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    title = path.stem
    if lines:
        first = re.sub(r"^#+\s*", "", lines[0]).strip()
        if first:
            title = first[:160]

    emails = sorted(set(re.findall(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text)))

    phones: list[str] = []
    for match in re.findall(r"\+?\d[\d\s().-]{6,}\d", text):
        cleaned = match.strip()
        digits = re.sub(r"\D", "", cleaned)

        # Reject common false positives such as dates and invoice fragments.
        is_date_like = bool(re.fullmatch(r"\d{4}-\d{2}-\d{2}", cleaned))
        is_year_code_like = bool(re.fullmatch(r"\d{4}-\d{3,}", cleaned))
        has_phone_shape = cleaned.startswith("+") or bool(re.search(r"\d[\s().-]+\d", cleaned))

        if 7 <= len(digits) <= 15 and has_phone_shape and not is_date_like and not is_year_code_like:
            phones.append(cleaned)

    phones = sorted(set(phones))

    dates: list[str] = []
    date_patterns = [
        r"\b\d{4}-\d{2}-\d{2}\b",
        r"\b\d{1,2}/\d{1,2}/\d{2,4}\b",
        r"\b\d{1,2}-\d{1,2}-\d{2,4}\b",
        r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b",
        r"\b\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{4}\b",
    ]
    for pattern in date_patterns:
        dates.extend(re.findall(pattern, text, flags=re.I))
    dates = sorted(set(dates))

    amounts: list[str] = []
    amount_patterns = [
        r"\bR\s?\d[\d,\s]*(?:\.\d{2})?\b",
        r"\bZAR\s?\d[\d,\s]*(?:\.\d{2})?\b",
        r"\$\s?\d[\d,\s]*(?:\.\d{2})?\b",
        r"\bEUR\s?\d[\d,\s]*(?:\.\d{2})?\b",
        r"\bGBP\s?\d[\d,\s]*(?:\.\d{2})?\b",
        r"\b\d[\d,\s]*(?:\.\d{2})?\s?(?:ZAR|USD|EUR|GBP|rand|rands|dollars|euros|pounds)\b",
    ]
    for pattern in amount_patterns:
        amounts.extend(re.findall(pattern, text, flags=re.I))
    amounts = sorted(set(x.strip() for x in amounts))

    possible_ids: list[str] = []
    id_labels = [
        "ID", "Ref", "Reference", "Invoice", "Invoice Number", "Claim",
        "Policy", "Member", "Membership", "Account", "Case", "File", "Patient"
    ]
    for line in lines:
        for label in id_labels:
            lower = line.lower()
            if lower.startswith(label.lower()):
                parts = re.split(r"[:#]", line, maxsplit=1)
                if len(parts) == 2:
                    value = parts[1].strip()
                    if value:
                        possible_ids.append(value)
        for match in re.findall(r"\b[A-Z]{2,}-\d{3,}(?:-\d+)?\b", line):
            possible_ids.append(match.strip())
        for match in re.findall(r"\bINV-\d{4}-\d{3,}\b", line, flags=re.I):
            possible_ids.append(match.strip())
    possible_ids = sorted(set(possible_ids))

    people_candidates = re.findall(r"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2}\b", text)
    stop_people = {
        "Medical Aid", "Billing Follow", "Billing Workflow", "Action Items",
        "Test Patient", "Dr Piet", "Dr Piet Muller", "South Africa"
    }
    people = sorted(set(x for x in people_candidates if x not in stop_people))[:30]

    organizations: list[str] = []
    org_keywords = ["practice", "medical aid", "clinic", "hospital", "company", "department", "scheme", "healthcare"]
    for line in lines:
        if any(k in line.lower() for k in org_keywords):
            organizations.append(line[:180])
    organizations = sorted(set(organizations))[:20]

    words = re.findall(r"\b[a-zA-Z][a-zA-Z-]{3,}\b", clean.lower())
    stopwords = {
        "this", "that", "with", "from", "have", "will", "should", "could", "would",
        "document", "information", "details", "before", "after", "about", "into",
        "their", "there", "where", "which", "what", "when", "then", "than", "also",
        "only", "does", "make", "clear", "person", "workflow", "support"
    }
    freq: dict[str, int] = {}
    for word in words:
        if word not in stopwords:
            freq[word] = freq.get(word, 0) + 1
    keywords = [word for word, _ in sorted(freq.items(), key=lambda kv: (-kv[1], kv[0]))[:20]]

    action_items: list[str] = []
    in_action_section = False
    for line in lines:
        lower = line.lower().rstrip(":")
        if lower in {"action items", "recommended next steps", "next steps", "to do", "todo"}:
            in_action_section = True
            continue

        if in_action_section and (line.startswith("- ") or line.startswith("* ")):
            action_items.append(line.lstrip("-* ").strip())
            continue

        if line.startswith("- ") or line.startswith("* "):
            action_items.append(line.lstrip("-* ").strip())
            continue

        if any(line.lower().startswith(prefix) for prefix in ["review ", "check ", "confirm ", "collect ", "track ", "escalate ", "submit ", "prepare ", "contact "]):
            action_items.append(line.strip())

    action_items = [x for x in action_items if x][:30]

    classification = _classify_document(path, text)
    inspection = _inspect_document(path, text)

    return {
        "source": str(path),
        "filename": path.name,
        "title": title,
        "document_type": classification["document_type"],
        "confidence": classification["confidence"],
        "sensitivity_flags": classification["sensitivity_flags"],
        "word_count": inspection["word_count"],
        "line_count": inspection["line_count"],
        "dates": dates,
        "emails": emails,
        "phone_numbers": phones,
        "amounts": amounts,
        "possible_ids": possible_ids,
        "people_candidates": people,
        "organization_contexts": organizations,
        "keywords": keywords,
        "action_items": action_items,
        "issues": inspection["issues"],
        "recommendations": inspection["recommendations"],
    }


def _save_json_artifact(kind: str, source: Path, payload: dict[str, Any]) -> Path:
    root = _documents_root()
    out_dir = root / kind
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"{_timestamp()}-{_safe_slug(source.stem)}.json"
    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    return path


@ToolRegistry.register("serena_documents_fields")
class SerenaDocumentsFieldsTool(_DocumentsBaseTool):
    tool_id = "serena_documents_fields"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Extract practical structured fields from a supported document.",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Document path."}
                },
                "required": ["path"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            path = _resolve_path(str(params.get("path") or ""))
            text = _extract_text(path)
            fields = _extract_structured_fields(path, text)
            out_path = _save_json_artifact("extracted", path, fields)

            lines = [
                "Serena document fields extracted",
                "",
                f"- Source: {path}",
                f"- JSON output: {out_path}",
                f"- Title: {fields['title']}",
                f"- Type: {fields['document_type']} ({fields['confidence']})",
                f"- Sensitivity flags: {', '.join(fields['sensitivity_flags']) or 'none'}",
                f"- Word count: {fields['word_count']}",
                "",
                "Extracted fields:",
                f"- Dates: {', '.join(fields['dates']) or 'none'}",
                f"- Emails: {', '.join(fields['emails']) or 'none'}",
                f"- Phone numbers: {', '.join(fields['phone_numbers']) or 'none'}",
                f"- Amounts: {', '.join(fields['amounts']) or 'none'}",
                f"- Possible IDs/references: {', '.join(fields['possible_ids']) or 'none'}",
                f"- People candidates: {', '.join(fields['people_candidates']) or 'none'}",
                f"- Keywords: {', '.join(fields['keywords'][:12]) or 'none'}",
                "",
                "Action items:",
            ]

            if fields["action_items"]:
                lines.extend(f"- {item}" for item in fields["action_items"][:15])
            else:
                lines.append("- none")

            return self._result(
                "\n".join(lines),
                metadata={"source": str(path), "json_output": str(out_path), **fields},
            )
        except Exception as exc:
            return self._result(f"Failed to extract document fields: {exc}", success=False)


@ToolRegistry.register("serena_documents_json_report")
class SerenaDocumentsJSONReportTool(_DocumentsBaseTool):
    tool_id = "serena_documents_json_report"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Create a structured JSON operator report for a supported document.",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Document path."}
                },
                "required": ["path"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            path = _resolve_path(str(params.get("path") or ""))
            text = _extract_text(path)
            fields = _extract_structured_fields(path, text)
            summary = _summarize_text(text, max_sentences=8)

            payload = {
                "report_type": "serena_document_operator_json_report",
                "created_at": _timestamp(),
                "summary": summary,
                "fields": fields,
                "operator_notes": [
                    "Original document was not modified.",
                    "Structured fields are pattern-based and should be reviewed for critical workflows.",
                    "For healthcare/legal/financial documents, Serena summarizes and flags but does not make final professional decisions.",
                ],
            }

            out_path = _save_json_artifact("reports", path, payload)

            return self._result(
                "Serena structured JSON document report created\n\n"
                f"- Source: {path}\n"
                f"- JSON report: {out_path}\n"
                f"- Type: {fields['document_type']}\n"
                f"- Sensitivity flags: {', '.join(fields['sensitivity_flags']) or 'none'}\n"
                f"- Dates found: {len(fields['dates'])}\n"
                f"- Emails found: {len(fields['emails'])}\n"
                f"- Phone numbers found: {len(fields['phone_numbers'])}\n"
                f"- Amounts found: {len(fields['amounts'])}\n"
                f"- Action items found: {len(fields['action_items'])}",
                metadata={
                    "source": str(path),
                    "json_report": str(out_path),
                    "document_type": fields["document_type"],
                    "sensitivity_flags": fields["sensitivity_flags"],
                    "field_counts": {
                        "dates": len(fields["dates"]),
                        "emails": len(fields["emails"]),
                        "phone_numbers": len(fields["phone_numbers"]),
                        "amounts": len(fields["amounts"]),
                        "possible_ids": len(fields["possible_ids"]),
                        "action_items": len(fields["action_items"]),
                    },
                },
            )
        except Exception as exc:
            return self._result(f"Failed to create structured JSON report: {exc}", success=False)


def _suggest_document_category(path: Path, text: str) -> str:
    classification = _classify_document(path, text)
    doc_type = classification["document_type"]

    mapping = {
        "clinical_or_health": "healthcare",
        "financial_or_billing": "billing-finance",
        "legal_or_compliance": "legal-compliance",
        "marketing_or_content": "marketing-content",
        "technical_or_project": "technical-projects",
        "cv_or_profile": "profiles-cvs",
        "general_document": "general",
    }

    return mapping.get(doc_type, "general")


def _safe_unique_target(folder: Path, filename: str) -> Path:
    target = folder / filename
    if not target.exists():
        return target

    stem = target.stem
    suffix = target.suffix
    counter = 2

    while True:
        candidate = folder / f"{stem}-{counter}{suffix}"
        if not candidate.exists():
            return candidate
        counter += 1


@ToolRegistry.register("serena_documents_plan_organize")
class SerenaDocumentsPlanOrganizeTool(_DocumentsBaseTool):
    tool_id = "serena_documents_plan_organize"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Plan how Serena would organize supported documents into categories without moving/copying files.",
            parameters={
                "type": "object",
                "properties": {
                    "folder": {"type": "string", "description": "Folder to scan."},
                    "recursive": {"type": "boolean", "description": "Scan recursively."},
                    "limit": {"type": "integer", "description": "Maximum files to plan."},
                },
                "required": ["folder"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            folder = _resolve_path(str(params.get("folder") or ""))
            recursive = bool(params.get("recursive", True))
            limit = int(params.get("limit") or 100)

            if not folder.is_dir():
                return self._result(f"Not a folder: {folder}", success=False)

            pattern = "**/*" if recursive else "*"
            files = [
                p for p in folder.glob(pattern)
                if p.is_file()
                and p.suffix.lower() in SUPPORTED_SUFFIXES
            ][:limit]

            plan: list[dict[str, Any]] = []

            for file in files:
                try:
                    text = _extract_text(file)
                    classification = _classify_document(file, text)
                    inspection = _inspect_document(file, text)
                    category = _suggest_document_category(file, text)
                    target_folder = _documents_root() / "library" / category
                    target_name = f"{_safe_slug(file.stem)}{file.suffix.lower()}"

                    plan.append(
                        {
                            "source": str(file),
                            "target_folder": str(target_folder),
                            "target_name": target_name,
                            "category": category,
                            "document_type": classification["document_type"],
                            "sensitivity_flags": classification["sensitivity_flags"],
                            "word_count": inspection["word_count"],
                            "issues": inspection["issues"],
                        }
                    )
                except Exception as exc:
                    plan.append(
                        {
                            "source": str(file),
                            "target_folder": "",
                            "target_name": "",
                            "category": "needs-review",
                            "document_type": "unreadable",
                            "sensitivity_flags": [],
                            "word_count": 0,
                            "issues": [f"extraction_failed: {exc}"],
                        }
                    )

            lines = [
                "Serena document organize plan",
                "",
                f"- Source folder: {folder}",
                f"- Recursive: {'yes' if recursive else 'no'}",
                f"- Supported documents planned: {len(plan)}",
                "",
                "Plan:",
            ]

            if not plan:
                lines.append("- none")
            else:
                for item in plan:
                    flags = ", ".join(item["sensitivity_flags"]) or "none"
                    lines.append(
                        f"- {Path(item['source']).name} -> {item['category']} | "
                        f"type {item['document_type']} | flags {flags} | issues {', '.join(item['issues']) or 'none'}"
                    )

            lines.extend([
                "",
                "Operator note:",
                "- This is a plan only. No files were moved or copied.",
                "- Use documents organize to copy into Serena's controlled library.",
                "- Use documents move only with explicit approval.",
            ])

            report_path = _save_json_artifact(
                "reports",
                folder,
                {
                    "report_type": "serena_documents_organize_plan",
                    "created_at": _timestamp(),
                    "source_folder": str(folder),
                    "recursive": recursive,
                    "plan": plan,
                },
            )

            return self._result(
                "\n".join(lines) + f"\n\nOrganize plan JSON saved: {report_path}",
                metadata={"source_folder": str(folder), "planned": len(plan), "plan_report": str(report_path), "plan": plan},
            )
        except Exception as exc:
            return self._result(f"Failed to plan document organization: {exc}", success=False)


@ToolRegistry.register("serena_documents_organize")
class SerenaDocumentsOrganizeTool(_DocumentsBaseTool):
    tool_id = "serena_documents_organize"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Copy supported documents into Serena's controlled library by detected category. Originals are preserved.",
            parameters={
                "type": "object",
                "properties": {
                    "folder": {"type": "string", "description": "Folder to organize."},
                    "recursive": {"type": "boolean", "description": "Scan recursively."},
                    "limit": {"type": "integer", "description": "Maximum files to organize."},
                },
                "required": ["folder"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            folder = _resolve_path(str(params.get("folder") or ""))
            recursive = bool(params.get("recursive", True))
            limit = int(params.get("limit") or 100)

            if not folder.is_dir():
                return self._result(f"Not a folder: {folder}", success=False)

            pattern = "**/*" if recursive else "*"
            files = [
                p for p in folder.glob(pattern)
                if p.is_file()
                and p.suffix.lower() in SUPPORTED_SUFFIXES
            ][:limit]

            copied: list[dict[str, Any]] = []
            skipped: list[dict[str, str]] = []

            for file in files:
                try:
                    text = _extract_text(file)
                    category = _suggest_document_category(file, text)
                    target_folder = _library_dir(category)
                    target = _safe_unique_target(target_folder, f"{_safe_slug(file.stem)}{file.suffix.lower()}")

                    shutil.copy2(file, target)

                    fields = _extract_structured_fields(target, text)
                    meta_path = target.with_suffix(target.suffix + ".json")
                    meta = {
                        "source": str(file),
                        "library_path": str(target),
                        "category": category,
                        "organized_at": _timestamp(),
                        "original_preserved": True,
                        "fields": fields,
                    }
                    meta_path.write_text(json.dumps(meta, indent=2), encoding="utf-8")

                    copied.append(
                        {
                            "source": str(file),
                            "target": str(target),
                            "metadata": str(meta_path),
                            "category": category,
                            "document_type": fields["document_type"],
                            "sensitivity_flags": fields["sensitivity_flags"],
                        }
                    )
                except Exception as exc:
                    skipped.append({"source": str(file), "error": str(exc)})

            lines = [
                "Serena documents organized into controlled library",
                "",
                f"- Source folder: {folder}",
                f"- Recursive: {'yes' if recursive else 'no'}",
                f"- Copied: {len(copied)}",
                f"- Skipped: {len(skipped)}",
                f"- Originals preserved: yes",
                "",
                "Copied documents:",
            ]

            if copied:
                for item in copied:
                    flags = ", ".join(item["sensitivity_flags"]) or "none"
                    lines.append(f"- {Path(item['source']).name} -> {item['category']} | {item['target']} | flags {flags}")
            else:
                lines.append("- none")

            lines.extend(["", "Skipped documents:"])
            if skipped:
                for item in skipped:
                    lines.append(f"- {item['source']}: {item['error']}")
            else:
                lines.append("- none")

            return self._result(
                "\n".join(lines),
                metadata={"source_folder": str(folder), "copied": copied, "skipped": skipped},
            )
        except Exception as exc:
            return self._result(f"Failed to organize documents: {exc}", success=False)


@ToolRegistry.register("serena_documents_copy")
class SerenaDocumentsCopyTool(_DocumentsBaseTool):
    tool_id = "serena_documents_copy"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Copy one document to a target folder without modifying the original.",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Source document path."},
                    "target_folder": {"type": "string", "description": "Target folder."},
                },
                "required": ["path", "target_folder"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            source = _resolve_path(str(params.get("path") or ""))
            target_folder = Path(str(params.get("target_folder") or ""))
            target_folder.mkdir(parents=True, exist_ok=True)

            if not source.is_file():
                return self._result(f"Not a file: {source}", success=False)

            target = _safe_unique_target(target_folder, source.name)
            shutil.copy2(source, target)

            return self._result(
                "Serena document copied\n\n"
                f"- Source: {source}\n"
                f"- Target: {target}\n"
                f"- Original preserved: yes",
                metadata={"source": str(source), "target": str(target), "original_preserved": True},
            )
        except Exception as exc:
            return self._result(f"Failed to copy document: {exc}", success=False)


@ToolRegistry.register("serena_documents_move")
class SerenaDocumentsMoveTool(_DocumentsBaseTool):
    tool_id = "serena_documents_move"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Move one document to a target folder only with explicit approval. Creates snapshot first.",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Source document path."},
                    "target_folder": {"type": "string", "description": "Target folder."},
                    "approved": {"type": "boolean", "description": "Required to move the original file."},
                },
                "required": ["path", "target_folder", "approved"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            source = _resolve_path(str(params.get("path") or ""))
            target_folder = Path(str(params.get("target_folder") or ""))
            approved = bool(params.get("approved", False))

            if not approved:
                return self._result(
                    "Document move blocked. Moving an original file requires explicit approval.",
                    success=False,
                    metadata={"source": str(source), "approved": False},
                )

            if not source.is_file():
                return self._result(f"Not a file: {source}", success=False)

            target_folder.mkdir(parents=True, exist_ok=True)
            snapshot = _snapshot_document(source, "before-move")
            target = _safe_unique_target(target_folder, source.name)

            shutil.move(str(source), str(target))

            return self._result(
                "Serena document moved with approval\n\n"
                f"- Source: {source}\n"
                f"- Target: {target}\n"
                f"- Snapshot before move: {snapshot}",
                metadata={"source": str(source), "target": str(target), "snapshot": str(snapshot), "approved": True},
            )
        except Exception as exc:
            return self._result(f"Failed to move document: {exc}", success=False)


@ToolRegistry.register("serena_documents_cleanup_candidates")
class SerenaDocumentsCleanupCandidatesTool(_DocumentsBaseTool):
    tool_id = "serena_documents_cleanup_candidates"

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name=self.tool_id,
            description="Find likely duplicate, empty, unsupported, or cleanup-candidate documents without deleting anything.",
            parameters={
                "type": "object",
                "properties": {
                    "folder": {"type": "string", "description": "Folder to scan."},
                    "recursive": {"type": "boolean", "description": "Scan recursively."},
                    "limit": {"type": "integer", "description": "Maximum files to scan."},
                },
                "required": ["folder"],
            },
            category="serena_documents",
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            folder = _resolve_path(str(params.get("folder") or ""))
            recursive = bool(params.get("recursive", True))
            limit = int(params.get("limit") or 500)

            if not folder.is_dir():
                return self._result(f"Not a folder: {folder}", success=False)

            pattern = "**/*" if recursive else "*"
            files = [p for p in folder.glob(pattern) if p.is_file()][:limit]

            by_name_size: dict[tuple[str, int], list[Path]] = {}
            unsupported: list[Path] = []
            empty: list[Path] = []

            for file in files:
                size = file.stat().st_size
                by_name_size.setdefault((file.name.lower(), size), []).append(file)

                if size == 0:
                    empty.append(file)

                if file.suffix.lower() not in SUPPORTED_SUFFIXES and not file.name.endswith(".json"):
                    unsupported.append(file)

            duplicates = [group for group in by_name_size.values() if len(group) > 1]

            lines = [
                "Serena document cleanup candidates",
                "",
                f"- Folder: {folder}",
                f"- Recursive: {'yes' if recursive else 'no'}",
                f"- Files scanned: {len(files)}",
                f"- Duplicate groups: {len(duplicates)}",
                f"- Empty files: {len(empty)}",
                f"- Unsupported files: {len(unsupported)}",
                "",
                "Duplicate candidates:",
            ]

            if duplicates:
                for group in duplicates[:20]:
                    lines.append("- group:")
                    for item in group:
                        lines.append(f"  - {item}")
            else:
                lines.append("- none")

            lines.extend(["", "Empty files:"])
            lines.extend(f"- {p}" for p in empty[:20]) if empty else lines.append("- none")

            lines.extend(["", "Unsupported files:"])
            lines.extend(f"- {p} | {p.suffix.lower()}" for p in unsupported[:20]) if unsupported else lines.append("- none")

            lines.extend([
                "",
                "Operator note:",
                "- No files were deleted.",
                "- Deletion/permanent cleanup is not part of Documents v1.",
                "- Use copy/organize first; use move only with explicit approval.",
            ])

            report_path = _save_json_artifact(
                "reports",
                folder,
                {
                    "report_type": "serena_documents_cleanup_candidates",
                    "created_at": _timestamp(),
                    "folder": str(folder),
                    "duplicates": [[str(p) for p in group] for group in duplicates],
                    "empty": [str(p) for p in empty],
                    "unsupported": [str(p) for p in unsupported],
                },
            )

            return self._result(
                "\n".join(lines) + f"\n\nCleanup candidate report saved: {report_path}",
                metadata={
                    "folder": str(folder),
                    "files_scanned": len(files),
                    "duplicate_groups": len(duplicates),
                    "empty_files": len(empty),
                    "unsupported_files": len(unsupported),
                    "report": str(report_path),
                },
            )
        except Exception as exc:
            return self._result(f"Failed to find cleanup candidates: {exc}", success=False)


__all__ = [
    "SerenaDocumentsStatusTool",
    "SerenaDocumentsIndexTool",
    "SerenaDocumentsReadTool",
    "SerenaDocumentsExtractTool",
    "SerenaDocumentsSummarizeTool",
    "SerenaDocumentsClassifyTool",
    "SerenaDocumentsInspectTool",
    "SerenaDocumentsReportTool",
    "SerenaDocumentsSnapshotsTool",
    "SerenaDocumentsAuditTool",
    "SerenaDocumentsPDFCheckTool",
    "SerenaDocumentsJSONReportTool",
    "SerenaDocumentsCleanupCandidatesTool",
    "SerenaDocumentsMoveTool",
    "SerenaDocumentsCopyTool",
    "SerenaDocumentsOrganizeTool",
    "SerenaDocumentsPlanOrganizeTool",
    "SerenaDocumentsFieldsTool",
    "SerenaDocumentsSnapshotTool",
    "SerenaDocumentsLibraryTool",
    "SerenaDocumentsImportTool",
]
